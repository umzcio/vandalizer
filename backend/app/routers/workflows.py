"""Workflow API routes."""

import asyncio
import base64
import csv
import io
import json
import logging
import re
import zipfile

from beanie import PydanticObjectId
from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, Request, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from app.dependencies import get_api_key_user, get_current_user
from app.models.user import User
from app.services import access_control
from app.services.access_control import get_authorized_search_set, get_authorized_workflow
from app.schemas.workflows import (
    AddStepRequest,
    AddTaskRequest,
    BatchStatusResponse,
    CreateTempDocumentsRequest,
    CreateWorkflowRequest,
    ReorderStepsRequest,
    RunWorkflowRequest,
    TestStepRequest,
    UpdateStepRequest,
    UpdateTaskRequest,
    UpdateValidationInputsRequest,
    UpdateValidationPlanRequest,
    UpdateWorkflowRequest,
    ValidateWorkflowResponse,
    ValidationInputsResponse,
    ValidationPlanResponse,
    WorkflowResponse,
    WorkflowStatusResponse,
)
from app.rate_limit import limiter
from app.services import workflow_service as svc
from app.services.user_lookup import resolve_author, resolve_authors

logger = logging.getLogger(__name__)

router = APIRouter()


async def _check_validation_input_documents_exist(uuids: list[str]) -> dict[str, bool]:
    """Return {uuid: exists} for the supplied document UUIDs.

    A soft-deleted doc is treated as gone — the UI should warn the user
    before they try to run the workflow with it.
    """
    from app.models.document import SmartDocument

    if not uuids:
        return {}
    unique = list({u for u in uuids if u})
    if not unique:
        return {}
    found = await SmartDocument.find(
        {"uuid": {"$in": unique}, "soft_deleted": {"$ne": True}}
    ).to_list()
    present = {d.uuid for d in found}
    return {u: (u in present) for u in unique}


async def _workflow_response_from_dict(wf: dict) -> WorkflowResponse:
    """Build a WorkflowResponse from a workflow dict, resolving the author."""
    creator_id = wf.get("created_by_user_id") or wf.get("user_id")
    created_by = await resolve_author(creator_id)
    return WorkflowResponse(**{**wf, "created_by": created_by})


def _csv_cell(value) -> str:
    """Format a value for a CSV cell, serializing complex types as JSON."""
    if isinstance(value, (dict, list)):
        return json.dumps(value, default=str)
    return str(value)


def _parse_structured_string(text: str):
    """Best-effort extraction of structured data from a free-form string output.

    Tries, in order: fenced JSON blocks, raw JSON, and GitHub-flavored markdown
    tables. Returns the parsed value (list/dict) or None if nothing matched.
    """
    if not isinstance(text, str):
        return None
    stripped = text.strip()
    if not stripped:
        return None

    # 1. Fenced code blocks (```json ... ``` or plain ``` ... ```)
    for match in re.finditer(r"```(?:json|JSON)?\s*\n?(.*?)```", text, re.DOTALL):
        candidate = match.group(1).strip()
        if not candidate:
            continue
        try:
            return json.loads(candidate)
        except (json.JSONDecodeError, ValueError):
            continue

    # 2. Raw JSON (already attempted upstream, but safe to retry)
    try:
        return json.loads(stripped)
    except (json.JSONDecodeError, ValueError):
        pass

    # 3. Markdown tables: a header row, a separator row of dashes, then body rows.
    lines = [ln.rstrip() for ln in text.splitlines()]
    for i in range(len(lines) - 1):
        header = lines[i].strip()
        sep = lines[i + 1].strip()
        if not (header.startswith("|") and sep.startswith("|")):
            continue
        if not re.fullmatch(r"\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?", sep):
            continue
        headers = [c.strip() for c in header.strip("|").split("|")]
        rows: list[dict] = []
        for body in lines[i + 2:]:
            body = body.strip()
            if not body.startswith("|"):
                break
            cells = [c.strip() for c in body.strip("|").split("|")]
            if len(cells) < len(headers):
                cells += [""] * (len(headers) - len(cells))
            rows.append({headers[j]: cells[j] for j in range(len(headers))})
        if rows:
            return rows

    return None


_INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
_INLINE_ITALIC = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_INLINE_CODE = re.compile(r"`(.+?)`")
_MD_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def _add_runs_with_formatting(paragraph, text: str) -> None:
    """Add runs to a docx paragraph, parsing inline **bold**, *italic*, `code`, [link](url)."""
    # Drop markdown link syntax — keep the visible text only (URLs in LLM output
    # are often hallucinated and useless in a Word doc).
    text = _MD_LINK.sub(r"\1", text)

    pos = 0
    # Combined scan: find the next inline marker and emit the preceding plain run.
    while pos < len(text):
        next_match = None
        next_kind = None
        for kind, pattern in (("bold", _INLINE_BOLD), ("code", _INLINE_CODE), ("italic", _INLINE_ITALIC)):
            m = pattern.search(text, pos)
            if m and (next_match is None or m.start() < next_match.start()):
                next_match = m
                next_kind = kind
        if next_match is None:
            paragraph.add_run(text[pos:])
            return
        if next_match.start() > pos:
            paragraph.add_run(text[pos:next_match.start()])
        run = paragraph.add_run(next_match.group(1))
        if next_kind == "bold":
            run.bold = True
        elif next_kind == "italic":
            run.italic = True
        elif next_kind == "code":
            run.font.name = "Consolas"
        pos = next_match.end()


# Separator row of a GitHub-flavored markdown table (e.g. ``|---|:--:|-|``).
# One-or-more dashes per column matches the GFM spec (and remark-gfm, which the
# UI renders with), so a model emitting ``|--|--|`` still yields a real table.
# Kept identical to pdf_service so DOCX and PDF detect the same tables.
_MD_TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$")
# Horizontal rule: 3+ repeats of the same -, * or _ (e.g. ``---``). Mirrors
# pdf_service so a standalone rule renders as a line, not literal text.
_MD_HR_RE = re.compile(r"^\s{0,3}([-*_])(\s*\1){2,}\s*$")


def _split_md_table_row(line: str) -> list[str]:
    """Split a markdown table row into trimmed cell strings."""
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    # Honor backslash-escaped pipes by stashing them behind a placeholder.
    line = line.replace(r"\|", "\x00")
    return [c.strip().replace("\x00", "|") for c in line.split("|")]


# Downloaded-table theme — matches the gold PDF header in pdf_service so Word and
# PDF downloads share the Vandalizer brand look (UI highlight color #eab308).
_DOCX_TABLE_HEADER_FILL = "EAB308"   # brand gold header fill
_DOCX_TABLE_STRIPE_FILL = "FDF9EB"   # faint gold tint for zebra striping
_DOCX_TABLE_BORDER = "E5E7EB"        # light-gray grid
_DOCX_HEADING_HEX = "191919"         # charcoal headings (override Word's blue)


def _shade_docx_cell(cell, fill_hex: str) -> None:
    """Set a table cell's background fill (python-docx has no high-level API)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_docx_table_borders(table, color_hex: str) -> None:
    """Apply a thin single-line grid border (color_hex) to every table edge."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")       # 4 eighths-of-a-point = 0.5pt
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        borders.append(el)
    table._tbl.tblPr.append(borders)


def _add_docx_hr(doc) -> None:
    """Render a markdown horizontal rule as an empty paragraph with a thin,
    light-gray bottom border (Word has no native ``<hr>``)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), _DOCX_TABLE_BORDER)
    borders.append(bottom)
    p_pr.append(borders)


def _apply_docx_table_theme(table) -> None:
    """Apply the Vandalizer gold table theme to a Word table.

    Gold (#eab308) header row with bold black text, faint gold zebra striping on
    body rows, and a light-gray grid. Replaces the built-in (blue) ``Light Grid
    Accent 1`` style so DOCX downloads match the gold-headed PDF tables.
    """
    table.style = "Table Grid"
    _set_docx_table_borders(table, _DOCX_TABLE_BORDER)
    rows = table.rows
    if not rows:
        return
    for cell in rows[0].cells:
        _shade_docx_cell(cell, _DOCX_TABLE_HEADER_FILL)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for idx, row in enumerate(rows[1:], start=1):
        if idx % 2 == 0:  # zebra-stripe every other body row
            for cell in row.cells:
                _shade_docx_cell(cell, _DOCX_TABLE_STRIPE_FILL)


def _add_markdown_table_to_docx(doc, headers: list[str], rows: list[list[str]]) -> None:
    """Render parsed markdown-table cells as a real Word table.

    Uses the shared gold table theme (see ``_apply_docx_table_theme``), with
    inline **bold**/*italic*/`code` honored inside every cell.
    """
    col_count = max(len(headers), max((len(r) for r in rows), default=0)) or 1
    headers = headers + [""] * (col_count - len(headers))

    table = doc.add_table(rows=1, cols=col_count)

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        _add_runs_with_formatting(hdr_cells[i].paragraphs[0], h)

    for row in rows:
        row = row + [""] * (col_count - len(row))
        cells = table.add_row().cells
        for i in range(col_count):
            _add_runs_with_formatting(cells[i].paragraphs[0], row[i])

    _apply_docx_table_theme(table)


def _markdown_to_docx(text: str):
    """Render a markdown-ish string into a python-docx Document.

    Handles ATX headings, unordered/ordered lists, GitHub-flavored tables,
    blank-line paragraphs, and inline bold/italic/code. Unknown syntax falls
    back to a plain paragraph.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    heading_color = RGBColor.from_string(_DOCX_HEADING_HEX)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    lines = (text or "").splitlines()
    n = len(lines)
    i = 0
    while i < n:
        line = lines[i].rstrip()
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # Markdown table — a header row followed by a dash-separator row. Consume
        # the header, separator, and all contiguous body rows into a real table.
        if "|" in line and i + 1 < n and _MD_TABLE_SEP_RE.match(lines[i + 1].rstrip()):
            headers = _split_md_table_row(line)
            i += 2  # skip header + separator
            rows: list[list[str]] = []
            while i < n:
                body = lines[i].rstrip()
                if not body.strip() or "|" not in body:
                    break
                rows.append(_split_md_table_row(body))
                i += 1
            _add_markdown_table_to_docx(doc, headers, rows)
            continue

        # Horizontal rule (``---`` / ``***`` / ``___``) → a thin border, not text.
        if _MD_HR_RE.match(line):
            _add_docx_hr(doc)
            i += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)), 4)
            h = doc.add_heading(heading.group(2).strip(), level=level)
            for run in h.runs:
                run.font.color.rgb = heading_color  # charcoal, not Word's blue
            i += 1
            continue

        bullet = re.match(r"^\s*[-*+]\s+(.+)$", line)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_formatting(p, bullet.group(1).strip())
            i += 1
            continue

        numbered = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            _add_runs_with_formatting(p, numbered.group(1).strip())
            i += 1
            continue

        p = doc.add_paragraph()
        _add_runs_with_formatting(p, line)
        i += 1

    return doc


def _data_to_docx_bytes(data) -> bytes:
    """Serialize workflow output (str/dict/list) to a .docx byte payload."""
    from docx import Document
    from docx.shared import Inches, Pt

    if isinstance(data, str):
        try:
            parsed = json.loads(data)
        except (json.JSONDecodeError, ValueError):
            parsed = None
        if isinstance(parsed, (dict, list)):
            data = parsed

    if isinstance(data, str):
        doc = _markdown_to_docx(data)
    elif isinstance(data, list) and data and isinstance(data[0], dict):
        doc = Document()
        for section in doc.sections:
            section.top_margin = section.bottom_margin = Inches(1)
            section.left_margin = section.right_margin = Inches(1)
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(11)
        headers = list(dict.fromkeys(k for row in data for k in row.keys()))
        table = doc.add_table(rows=1, cols=len(headers))
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = str(h)
        for row in data:
            cells = table.add_row().cells
            for i, h in enumerate(headers):
                val = row.get(h, "")
                cells[i].text = json.dumps(val, default=str) if isinstance(val, (dict, list)) else str(val if val is not None else "")
        _apply_docx_table_theme(table)
    elif isinstance(data, dict):
        doc = Document()
        for section in doc.sections:
            section.top_margin = section.bottom_margin = Inches(1)
            section.left_margin = section.right_margin = Inches(1)
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(11)
        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "Field"
        hdr[1].text = "Value"
        for k, v in data.items():
            cells = table.add_row().cells
            cells[0].text = str(k)
            cells[1].text = json.dumps(v, default=str) if isinstance(v, (dict, list)) else str(v)
        _apply_docx_table_theme(table)
    elif isinstance(data, list):
        doc = _markdown_to_docx("\n".join(f"- {item}" for item in data))
    else:
        doc = _markdown_to_docx("" if data is None else str(data))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def _authorize_documents(document_uuids: list[str], user: User) -> list[str]:
    team_access = await access_control.get_team_access_context(user)
    authorized: list[str] = []
    for doc_uuid in document_uuids:
        doc = await access_control.get_authorized_document(
            doc_uuid,
            user,
            team_access=team_access,
            allow_admin=True,
        )
        if not doc:
            raise HTTPException(status_code=404, detail=f"Document not found: {doc_uuid}")
        authorized.append(doc.uuid)
    return authorized


# ---------------------------------------------------------------------------
# Workflow CRUD
# ---------------------------------------------------------------------------

@router.post("", response_model=WorkflowResponse)
async def create_workflow(req: CreateWorkflowRequest, user: User = Depends(get_current_user)):
    team_id = str(user.current_team) if user.current_team else None
    wf = await svc.create_workflow(req.name, user.user_id, req.description, team_id=team_id)
    created_by = await resolve_author(wf.created_by_user_id or wf.user_id)
    return WorkflowResponse(
        id=str(wf.id), name=wf.name, description=wf.description,
        user_id=wf.user_id, team_id=wf.team_id, num_executions=wf.num_executions,
        can_manage=True,  # creator can always manage
        created_by=created_by,
    )


@router.get("", response_model=list[WorkflowResponse])
async def list_workflows(
    skip: int = Query(default=0, ge=0),
    limit: int = Query(default=100, ge=1, le=500),
    scope: str | None = Query(default=None),
    search: str | None = Query(default=None),
    user: User = Depends(get_current_user),
):
    workflows = await svc.list_workflows(user=user, skip=skip, limit=limit, scope=scope, search=search)
    author_map = await resolve_authors(
        (wf.created_by_user_id or wf.user_id) for wf in workflows
    )
    # One team-access lookup powers can_manage for every workflow in the page.
    team_access = await access_control.get_team_access_context(user)
    return [
        WorkflowResponse(
            id=str(wf.id), name=wf.name, description=wf.description,
            user_id=wf.user_id, team_id=wf.team_id, num_executions=wf.num_executions,
            can_manage=access_control.can_manage_workflow(wf, user, team_access),
            created_by=author_map.get(wf.created_by_user_id or wf.user_id),
        )
        for wf in workflows
    ]


@router.get("/status", response_model=WorkflowStatusResponse)
async def get_workflow_status(session_id: str, user: User = Depends(get_current_user)):
    status = await svc.get_workflow_status(session_id, user=user)
    if not status:
        raise HTTPException(status_code=404, detail="Workflow result not found")
    return WorkflowStatusResponse(**status)


@router.get("/batch-status", response_model=BatchStatusResponse)
async def get_batch_status(batch_id: str, user: User = Depends(get_current_user)):
    status = await svc.get_batch_status(batch_id, user=user)
    if not status:
        raise HTTPException(status_code=404, detail="Batch not found")
    return status


@router.get("/status/stream")
async def stream_workflow_status(session_id: str, user: User = Depends(get_current_user)):
    """SSE endpoint that streams workflow status updates until completion."""
    initial_status = await svc.get_workflow_status(session_id, user=user)
    if not initial_status:
        raise HTTPException(status_code=404, detail="Workflow result not found")

    async def event_generator():
        last_json = ""
        not_found_retries = 0
        while True:
            status = await svc.get_workflow_status(session_id, user=user)
            if not status:
                not_found_retries += 1
                # Allow a few retries for the workflow result to appear in the DB
                if not_found_retries > 10:
                    yield f"data: {json.dumps({'error': 'not_found'})}\n\n"
                    return
                await asyncio.sleep(1.0)
                continue

            current_json = json.dumps(status, default=str)
            # Only send if something changed
            if current_json != last_json:
                last_json = current_json
                yield f"data: {current_json}\n\n"

            if status.get("status") in ("completed", "error", "failed"):
                return

            await asyncio.sleep(1.5)

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@router.get("/steps/test/{task_id}")
async def poll_step_test(task_id: str, user: User = Depends(get_current_user)):
    return svc.get_test_status(task_id)


def _step_output_value(step_payload):
    """Extract the displayable value from a stored step output payload.

    Each entry in WorkflowResult.steps_output is the full node return dict
    ({"output": ..., "input": ..., "step_name": ..., "formatted_output": ...}).
    """
    if not isinstance(step_payload, dict):
        return step_payload
    return step_payload.get("formatted_output") or step_payload.get("output")


def _zip_member_for_step(step_name: str, value):
    """Return (filename, bytes) for one step's output, for inclusion in a multi-output zip."""
    safe_name = re.sub(r"[^A-Za-z0-9_\-]+", "_", step_name).strip("_") or "step"

    if isinstance(value, dict) and value.get("type") == "file_download":
        filename = value.get("filename") or f"{safe_name}.bin"
        return filename, base64.b64decode(value.get("data_b64", ""))

    if isinstance(value, (dict, list)):
        return f"{safe_name}.json", json.dumps(value, indent=2, default=str).encode()

    if value is None:
        return f"{safe_name}.txt", b""

    return f"{safe_name}.txt", str(value).encode()


@router.get("/download")
async def download_results(
    session_id: str,
    format: str = "json",
    parse_structured: bool = False,
    user: User = Depends(get_current_user),
):
    """Download workflow results in specified format.

    If the workflow has multiple steps marked as deliverables (is_output), the
    response is a ZIP bundle containing one file per marked step. With 0 or 1
    marked steps the single-output formatting paths below apply.
    """
    status = await svc.get_workflow_status(session_id, user=user)
    if not status:
        raise HTTPException(status_code=404, detail="Workflow result not found")

    # Build a base filename unique per session. Browsers cap auto-suffixing of
    # duplicate downloads at ~5; past that, the same Content-Disposition name
    # causes older files to be overwritten. Embedding the session id in the
    # name guarantees uniqueness across manual runs.
    workflow_name = status.get("workflow_name")
    document_title = status.get("document_title")
    name_parts: list[str] = []
    if workflow_name:
        name_parts.append(workflow_name)
    else:
        name_parts.append("results")
    if document_title:
        doc_stem = document_title.rsplit(".", 1)[0] if "." in document_title else document_title
        name_parts.append(doc_stem)
    name_parts.append(session_id[:8])
    raw_base = "-".join(name_parts)
    base_filename = "".join(c if c.isalnum() or c in " _-." else "_" for c in raw_base).strip() or f"results-{session_id[:8]}"

    final_output = status.get("final_output", {})
    steps_output = status.get("steps_output", {}) or {}
    output_step_names = [n for n in (status.get("output_step_names") or []) if n in steps_output]

    if len(output_step_names) >= 2:
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            seen: dict[str, int] = {}
            for name in output_step_names:
                value = _step_output_value(steps_output.get(name))
                filename, payload = _zip_member_for_step(name, value)
                # Disambiguate duplicate filenames.
                if filename in seen:
                    seen[filename] += 1
                    stem, _, ext = filename.rpartition(".")
                    filename = f"{stem}_{seen[filename]}.{ext}" if ext else f"{filename}_{seen[filename]}"
                else:
                    seen[filename] = 0
                zf.writestr(filename, payload)
        zip_buf.seek(0)
        return StreamingResponse(
            zip_buf,
            media_type="application/zip",
            headers={"Content-Disposition": f'attachment; filename="{base_filename}.zip"'},
        )

    if len(output_step_names) == 1:
        output_data = _step_output_value(steps_output.get(output_step_names[0]))
    else:
        output_data = final_output.get("output", "") if isinstance(final_output, dict) else final_output

    # Check for file_download type (e.g., from DataExport or DocumentRenderer)
    if isinstance(output_data, dict) and output_data.get("type") == "file_download":
        file_bytes = base64.b64decode(output_data["data_b64"])
        media_type_map = {"pdf": "application/pdf", "csv": "text/csv", "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "json": "application/json", "zip": "application/zip"}
        media_type = media_type_map.get(output_data.get("file_type", ""), "application/octet-stream")
        return StreamingResponse(
            io.BytesIO(file_bytes),
            media_type=media_type,
            headers={"Content-Disposition": f'attachment; filename="{output_data.get("filename", "output")}"'},
        )

    if format == "csv":
        buf = io.StringIO()
        writer = csv.writer(buf)
        # Parse JSON strings so fields land in separate columns
        data = output_data
        if isinstance(data, str):
            try:
                data = json.loads(data)
            except (json.JSONDecodeError, ValueError):
                if parse_structured:
                    parsed = _parse_structured_string(data)
                    if parsed is not None:
                        data = parsed
        if isinstance(data, list):
            if data and isinstance(data[0], dict):
                # Collect keys from ALL items so no columns are missing
                headers = list(dict.fromkeys(k for row in data for k in row.keys()))
                writer.writerow(headers)
                for row in data:
                    writer.writerow([_csv_cell(row.get(h, "")) for h in headers])
            else:
                writer.writerow(["Value"])
                for item in data:
                    writer.writerow([str(item)])
        elif isinstance(data, dict):
            # Transpose to Field/Value rows instead of one wide row
            writer.writerow(["Field", "Value"])
            for k, v in data.items():
                writer.writerow([str(k), _csv_cell(v)])
        else:
            writer.writerow(["Output"])
            text = str(data)
            for line in text.split("\n"):
                if line.strip():
                    writer.writerow([line])
        return StreamingResponse(
            io.BytesIO(buf.getvalue().encode()),
            media_type="text/csv",
            headers={"Content-Disposition": f'attachment; filename="{base_filename}.csv"'},
        )

    if format == "text":
        if isinstance(output_data, str):
            text = output_data
        elif isinstance(output_data, dict):
            parts = []
            for k, v in output_data.items():
                parts.append(f"{k}: {v}")
            text = "\n".join(parts)
        elif isinstance(output_data, list):
            text = "\n".join(str(item) for item in output_data)
        else:
            text = str(output_data)
        return StreamingResponse(
            io.BytesIO(text.encode()),
            media_type="text/plain",
            headers={"Content-Disposition": f'attachment; filename="{base_filename}.txt"'},
        )

    if format == "pdf":
        from app.services.pdf_service import render_workflow_pdf

        pdf_bytes = render_workflow_pdf(output_data, title="Workflow Results")
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{base_filename}.pdf"'},
        )

    if format == "docx":
        docx_bytes = _data_to_docx_bytes(output_data)
        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{base_filename}.docx"'},
        )

    # Default: JSON
    json_bytes = json.dumps(output_data, indent=2, default=str).encode()
    return StreamingResponse(
        io.BytesIO(json_bytes),
        media_type="application/json",
        headers={"Content-Disposition": f'attachment; filename="{base_filename}.json"'},
    )


class SaveOutputToFolderRequest(BaseModel):
    folder_uuid: str
    format: str = "pdf"  # pdf | markdown | csv | json | text
    file_name: str | None = None


@router.post("/sessions/{session_id}/save-to-folder")
async def save_session_output_to_folder(
    session_id: str,
    req: SaveOutputToFolderRequest,
    user: User = Depends(get_current_user),
):
    """Save a workflow run's output as a SmartDocument in the chosen SmartFolder.

    Writes the rendered output to disk (PDF/Markdown/CSV/JSON/text) and inserts
    a SmartDocument record so the file shows up in the user's file structure.
    """
    from app.models.workflow import WorkflowResult
    from app.services.access_control import get_authorized_folder
    from app.services.output_handlers import save_results_to_folder

    valid_formats = {"pdf", "markdown", "csv", "json", "text"}
    if req.format not in valid_formats:
        raise HTTPException(status_code=400, detail=f"Invalid format. Use one of: {sorted(valid_formats)}")

    result = await WorkflowResult.find_one(WorkflowResult.session_id == session_id)
    if not result or not result.workflow:
        raise HTTPException(status_code=404, detail="Workflow result not found")
    wf = await get_authorized_workflow(str(result.workflow), user)
    if not wf:
        raise HTTPException(status_code=404, detail="Workflow result not found")

    folder = await get_authorized_folder(req.folder_uuid, user, manage=True)
    if not folder:
        raise HTTPException(status_code=404, detail="Folder not found")

    storage_cfg: dict = {
        "destination_folder": req.folder_uuid,
        "format": req.format,
        "on_rerun": "new",
        "actor_user_id": user.user_id,
    }
    if req.file_name:
        safe = "".join(c if c.isalnum() or c in " _-." else "_" for c in req.file_name).strip()
        if safe:
            storage_cfg["file_naming"] = safe.rsplit(".", 1)[0] if "." in safe else safe

    result_doc = result.model_dump(by_alias=True)
    result_doc["_id"] = result.id
    result_doc["workflow"] = result.workflow

    try:
        file_path = await asyncio.to_thread(save_results_to_folder, result_doc, storage_cfg)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception:
        logger.exception("Failed to save workflow output to folder")
        raise HTTPException(status_code=500, detail="Failed to save output")

    return {"ok": True, "folder_uuid": req.folder_uuid, "file_path": file_path}


@router.get("/{workflow_id}/export")
async def export_workflow(workflow_id: str, user: User = Depends(get_current_user)):
    """Download workflow definition as a shareable JSON file."""
    wf = await get_authorized_workflow(workflow_id, user)
    if not wf:
        raise HTTPException(status_code=404, detail="Workflow not found")

    from app.services import export_import_service as eis

    try:
        data = await eis.export_workflow(workflow_id, user.email or user.user_id)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))

    json_bytes = json.dumps(data, indent=2, default=str).encode()
    safe_name = "".join(c if c.isalnum() or c in " _-" else "_" for c in (data["items"][0]["name"] or "workflow")).strip() or "workflow"
    return StreamingResponse(
        io.BytesIO(json_bytes),
        media_type="application/json",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}.vandalizer.json"'},
    )


@router.post("/import")
async def import_workflow(
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
):
    """Import a workflow from an exported JSON file."""
    from app.services import export_import_service as eis

    content = await file.read()
    try:
        data = json.loads(content)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid JSON file")

    try:
        team_id = str(user.current_team) if user.current_team else None
        wf = await eis.import_workflow(data, user.user_id, team_id=team_id)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    return await _workflow_response_from_dict(wf)


@router.post("/{workflow_id}/import", response_model=WorkflowResponse)
async def import_into_workflow(
    workflow_id: str,
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
):
    """Replace the contents of an existing workflow with imported JSON data.

    The workflow's id, name, and description are preserved; steps, tasks,
    and configs are replaced.
    """
    from app.services import export_import_service as eis

    target = await get_authorized_workflow(workflow_id, user, manage=True)
    if not target:
        raise HTTPException(status_code=404, detail="Workflow not found")

    content = await file.read()
    try:
        data = json.loads(content)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid JSON file")

    try:
        team_id = str(user.current_team) if user.current_team else None
        wf = await eis.import_into_workflow(workflow_id, data, user.user_id, team_id=team_id)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    return await _workflow_response_from_dict(wf)


@router.get("/{workflow_id}", response_model=WorkflowResponse)
async def get_workflow(
    workflow_id: str,
    share_token: str | None = Query(default=None),
    user: User = Depends(get_current_user),
):
    wf = await svc.get_workflow(workflow_id, user=user, share_token=share_token)
    if not wf:
        raise HTTPException(status_code=404, detail="Workflow not found")
    return await _workflow_response_from_dict(wf)


@router.post("/{workflow_id}/share-token")
async def mint_workflow_share_token(workflow_id: str, user: User = Depends(get_current_user)):
    """Mint (or return existing) view-only share token for a workflow.

    Manager-level access required: owners and team owner/admin can issue
    share links. Anyone holding the token can view and duplicate the
    workflow but cannot edit the original.
    """
    import secrets

    wf = await get_authorized_workflow(workflow_id, user, manage=True)
    if not wf:
        raise HTTPException(status_code=404, detail="Workflow not found")
    if not wf.share_token:
        wf.share_token = secrets.token_urlsafe(32)
        await wf.save()
    return {"share_token": wf.share_token}


@router.patch("/{workflow_id}", response_model=WorkflowResponse)
async def update_workflow(workflow_id: str, req: UpdateWorkflowRequest, user: User = Depends(get_current_user)):
    wf = await svc.update_workflow(
        workflow_id, user=user, name=req.name, description=req.description,
        input_config=req.input_config, output_config=req.output_config,
    )
    if not wf:
        raise HTTPException(status_code=404, detail="Workflow not found")
    # Flag stale verification if this workflow was verified
    from app.services.verification_service import check_and_flag_stale_verification
    await check_and_flag_stale_verification("workflow", str(wf.id))
    created_by = await resolve_author(wf.created_by_user_id or wf.user_id)
    return WorkflowResponse(
        id=str(wf.id), name=wf.name, description=wf.description,
        user_id=wf.user_id, team_id=wf.team_id, num_executions=wf.num_executions,
        input_config=wf.input_config,
        output_config=wf.output_config,
        can_manage=True,  # update already enforced manage authorization
        created_by=created_by,
    )


@router.delete("/{workflow_id}")
async def delete_workflow(workflow_id: str, user: User = Depends(get_current_user)):
    ok = await svc.delete_workflow(workflow_id, user=user)
    if not ok:
        raise HTTPException(status_code=404, detail="Workflow not found")
    return {"ok": True}


@router.post("/{workflow_id}/duplicate", response_model=WorkflowResponse)
async def duplicate_workflow(
    workflow_id: str,
    share_token: str | None = Query(default=None),
    user: User = Depends(get_current_user),
):
    team_id = str(user.current_team) if user.current_team else None
    wf = await svc.duplicate_workflow(
        workflow_id,
        user=user,
        user_id=user.user_id,
        team_id=team_id,
        share_token=share_token,
    )
    if not wf:
        raise HTTPException(status_code=404, detail="Workflow not found")
    return await _workflow_response_from_dict(wf)


@router.delete("/{workflow_id}/team", response_model=WorkflowResponse)
async def remove_workflow_from_team(workflow_id: str, user: User = Depends(get_current_user)):
    """Remove a workflow from its team library without deleting it.

    The workflow stays owned by its creator (``user_id``) and disappears from
    every other team member's view. Allowed for the creator or a team
    owner/admin (same set as ``can_manage_workflow``).
    """
    try:
        wf = await svc.remove_workflow_from_team(workflow_id, user=user)
    except svc.WorkflowNotInTeam:
        raise HTTPException(status_code=400, detail="Workflow is not in a team")
    if not wf:
        raise HTTPException(status_code=404, detail="Workflow not found")
    created_by = await resolve_author(wf.created_by_user_id or wf.user_id)
    return WorkflowResponse(
        id=str(wf.id),
        name=wf.name,
        description=wf.description,
        user_id=wf.user_id,
        team_id=wf.team_id,
        num_executions=wf.num_executions,
        input_config=wf.input_config,
        output_config=wf.output_config,
        can_manage=True,  # caller just managed it; trivially true
        created_by=created_by,
    )


# ---------------------------------------------------------------------------
# Steps
# ---------------------------------------------------------------------------

@router.post("/{workflow_id}/steps")
async def add_step(workflow_id: str, req: AddStepRequest, user: User = Depends(get_current_user)):
    step = await svc.add_step(workflow_id, req.name, user=user, data=req.data, is_output=req.is_output)
    if not step:
        raise HTTPException(status_code=404, detail="Workflow not found")
    return step


@router.patch("/steps/{step_id}")
async def update_step(step_id: str, req: UpdateStepRequest, user: User = Depends(get_current_user)):
    step = await svc.update_step(step_id, user=user, name=req.name, data=req.data, is_output=req.is_output)
    if not step:
        raise HTTPException(status_code=404, detail="Step not found")
    # Flag stale verification on parent workflow
    from app.services.verification_service import check_and_flag_stale_verification
    wf = await svc._get_workflow_for_step(PydanticObjectId(step_id))
    if wf:
        await check_and_flag_stale_verification("workflow", str(wf.id))
    return step


@router.delete("/steps/{step_id}")
async def delete_step(step_id: str, user: User = Depends(get_current_user)):
    ok = await svc.delete_step(step_id, user=user)
    if not ok:
        raise HTTPException(status_code=404, detail="Step not found")
    return {"ok": True}


# ---------------------------------------------------------------------------
# Tasks
# ---------------------------------------------------------------------------

@router.post("/steps/{step_id}/tasks")
async def add_task(step_id: str, req: AddTaskRequest, user: User = Depends(get_current_user)):
    task = await svc.add_task(step_id, req.name, user=user, data=req.data)
    if not task:
        status_code, detail = await _diagnose_step_mutation_failure(step_id, user)
        raise HTTPException(status_code=status_code, detail=detail)
    return task


async def _diagnose_step_mutation_failure(step_id: str, user: User) -> tuple[int, str]:
    """Why did a step-scoped mutation return None?

    The service collapses "step missing", "orphan workflow", and "user lacks
    permission" into the same `None` return. This walks the same lookups to
    pick a clearer status + detail for the client.
    """
    from app.models.workflow import WorkflowStep

    try:
        step = await WorkflowStep.get(PydanticObjectId(step_id))
    except Exception:
        return 404, "Step not found"
    if not step:
        return 404, "Step not found"
    parent = await svc._get_workflow_for_step(step.id)
    if not parent:
        return 404, "Step's workflow not found"
    team_access = await access_control.get_team_access_context(user)
    if not access_control.can_manage_workflow(parent, user, team_access):
        return 403, "You don't have permission to edit this workflow"
    return 500, "Failed to update step"


@router.patch("/tasks/{task_id}")
async def update_task(task_id: str, req: UpdateTaskRequest, user: User = Depends(get_current_user)):
    task = await svc.update_task(task_id, user=user, name=req.name, data=req.data)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    # Flag stale verification on parent workflow
    from app.services.verification_service import check_and_flag_stale_verification
    wf = await svc._get_workflow_for_task(PydanticObjectId(task_id))
    if wf:
        await check_and_flag_stale_verification("workflow", str(wf.id))
    return task


@router.delete("/tasks/{task_id}")
async def delete_task(task_id: str, user: User = Depends(get_current_user)):
    ok = await svc.delete_task(task_id, user=user)
    if not ok:
        raise HTTPException(status_code=404, detail="Task not found")
    return {"ok": True}


# ---------------------------------------------------------------------------
# Execution
# ---------------------------------------------------------------------------

@router.post("/{workflow_id}/run")
@limiter.limit("20/minute")
async def run_workflow(request: Request, workflow_id: str, req: RunWorkflowRequest, user: User = Depends(get_current_user)):
    from app.models.activity import ActivityType
    from app.services import activity_service
    from beanie import PydanticObjectId

    # Authorize and look up workflow name and step count for activity
    wf = await get_authorized_workflow(workflow_id, user)
    if not wf:
        raise HTTPException(status_code=404, detail="Workflow not found")
    document_uuids = await _authorize_documents(req.document_uuids, user)
    initial_title = wf.name if wf else "Workflow Run"
    # steps count excludes the trigger step
    num_steps = max(0, len(wf.steps) - 1) if wf and wf.steps else 0

    activity = await activity_service.activity_start(
        type=ActivityType.WORKFLOW_RUN,
        title=initial_title,
        user_id=user.user_id,
        team_id=str(user.current_team) if user.current_team else None,
        workflow=PydanticObjectId(workflow_id),
        steps_total=num_steps,
    )

    try:
        if req.batch_mode and len(document_uuids) > 1:
            batch_id = await svc.run_workflow_batch(
                workflow_id, document_uuids, user.user_id, req.model,
                activity_id=str(activity.id),
                user=user,
            )
            return {"batch_id": batch_id, "activity_id": str(activity.id)}
        else:
            session_id = await svc.run_workflow(
                workflow_id, document_uuids, user.user_id, req.model,
                activity_id=str(activity.id),
                user=user,
            )
            activity.workflow_session_id = session_id
            await activity.save()
            return {"session_id": session_id, "activity_id": str(activity.id)}
    except ValueError as e:
        from app.models.activity import ActivityStatus
        await activity_service.activity_finish(activity.id, ActivityStatus.FAILED, error=str(e))
        raise HTTPException(status_code=404, detail=str(e))


@router.post("/sessions/{session_id}/cancel")
async def cancel_workflow_run(session_id: str, user: User = Depends(get_current_user)):
    """Stop an in-flight workflow run (single-run sessions)."""
    result = await svc.cancel_workflow(session_id, user)
    if result is None:
        raise HTTPException(status_code=404, detail="Workflow run not found")
    return result


@router.post("/steps/test")
@limiter.limit("20/minute")
async def test_step(request: Request, req: TestStepRequest, user: User = Depends(get_current_user)):
    document_uuids = await _authorize_documents(req.document_uuids, user)
    if req.task_name == "Extraction":
        search_set_uuid = (req.task_data or {}).get("search_set_uuid")
        if search_set_uuid:
            ss = await get_authorized_search_set(search_set_uuid, user)
            if not ss:
                raise HTTPException(status_code=404, detail="Search set not found")
    task_id = await svc.test_step(
        req.task_name, req.task_data, document_uuids, user.user_id, req.model
    )
    return {"task_id": task_id}


@router.post("/{workflow_id}/reorder-steps")
async def reorder_steps(workflow_id: str, req: ReorderStepsRequest, user: User = Depends(get_current_user)):
    ok = await svc.reorder_steps(workflow_id, req.step_ids, user=user)
    if not ok:
        raise HTTPException(status_code=400, detail="Invalid step IDs or workflow not found")
    return {"ok": True}


@router.get("/{workflow_id}/history")
async def get_workflow_history(
    workflow_id: str, limit: int = 50, user: User = Depends(get_current_user),
):
    """List the current user's past runs of this workflow."""
    wf = await get_authorized_workflow(workflow_id, user)
    if not wf:
        raise HTTPException(status_code=404, detail="Workflow not found")
    from app.models.activity import ActivityEvent
    events = (
        await ActivityEvent.find(
            ActivityEvent.workflow == wf.id,
            ActivityEvent.user_id == user.user_id,
            ActivityEvent.type == "workflow_run",
        )
        .sort("-started_at")
        .limit(limit)
        .to_list()
    )
    return {
        "runs": [
            {
                "id": str(ev.id),
                "status": ev.status,
                "started_at": ev.started_at.isoformat() if ev.started_at else None,
                "finished_at": ev.finished_at.isoformat() if ev.finished_at else None,
                "duration_ms": ev.duration_ms,
                "error": ev.error or "",
                "tokens_input": ev.tokens_input,
                "tokens_output": ev.tokens_output,
                "documents_touched": ev.documents_touched,
                "steps_completed": ev.steps_completed,
                "steps_total": ev.steps_total,
                "session_id": ev.workflow_session_id,
                "result_snapshot": ev.result_snapshot or {},
            }
            for ev in events
        ],
    }


@router.get("/{workflow_id}/quality-history")
async def get_workflow_quality_history(
    workflow_id: str, limit: int = 50, user: User = Depends(get_current_user),
):
    wf = await get_authorized_workflow(workflow_id, user)
    if not wf:
        raise HTTPException(status_code=404, detail="Workflow not found")
    from app.services.quality_service import get_quality_history
    return {"runs": await get_quality_history("workflow", workflow_id, limit)}


@router.get("/{workflow_id}/quality-sparkline")
async def get_workflow_quality_sparkline(
    workflow_id: str, limit: int = 10, user: User = Depends(get_current_user),
):
    """Return compact score history for sparkline visualization."""
    wf = await get_authorized_workflow(workflow_id, user)
    if not wf:
        raise HTTPException(status_code=404, detail="Workflow not found")
    from app.services.quality_service import get_quality_history
    runs = await get_quality_history("workflow", workflow_id, limit)
    scores = [{"score": r["score"], "created_at": r["created_at"]} for r in reversed(runs)]
    return {"scores": scores}


@router.post("/{workflow_id}/improvement-suggestions")
@limiter.limit("5/minute")
async def get_workflow_suggestions(
    request: Request,
    workflow_id: str, user: User = Depends(get_current_user),
):
    """Use LLM to suggest improvements based on the latest validation run."""
    wf = await get_authorized_workflow(workflow_id, user)
    if not wf:
        raise HTTPException(status_code=404, detail="Workflow not found")

    from app.services.quality_service import get_latest_validation, generate_improvement_suggestions

    latest = await get_latest_validation("workflow", workflow_id)
    if not latest:
        raise HTTPException(status_code=404, detail="No validation runs found for this workflow")
    result_snapshot = latest.get("result_snapshot", latest)
    suggestions = await generate_improvement_suggestions("workflow", workflow_id, result_snapshot)
    return {"suggestions": suggestions}


class ImprovePromptRequest(BaseModel):
    prompt: str
    input_source: str | None = None
    prev_step_name: str | None = None
    sample_input: str | None = None


@router.post("/improve-prompt")
@limiter.limit("10/minute")
async def improve_prompt_endpoint(
    request: Request,
    body: ImprovePromptRequest,
    user: User = Depends(get_current_user),
):
    """LLM-suggested rewrite of a Prompt task's prompt, with rationale."""
    if not body.prompt.strip():
        raise HTTPException(status_code=400, detail="Prompt is required")
    from app.services.prompt_improvement_service import improve_prompt
    try:
        return await improve_prompt(
            prompt=body.prompt,
            input_source=body.input_source,
            prev_step_name=body.prev_step_name,
            sample_input=body.sample_input,
        )
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Failed to generate suggestion: {exc}")


@router.get("/{workflow_id}/quality-status")
async def get_workflow_quality_status(
    workflow_id: str, user: User = Depends(get_current_user),
):
    """Return quality status for Quality Pulse card (mirrors extraction quality-status)."""
    import hashlib
    import datetime as _dt
    from app.models.verification import VerifiedItemMetadata
    from app.services.quality_service import get_latest_validation

    wf = await get_authorized_workflow(workflow_id, user)
    if not wf:
        raise HTTPException(status_code=404, detail="Workflow not found")

    meta = await VerifiedItemMetadata.find_one(
        VerifiedItemMetadata.item_kind == "workflow",
        VerifiedItemMetadata.item_id == workflow_id,
    )
    latest = await get_latest_validation("workflow", workflow_id)

    if not latest and not meta:
        return {"status": "unvalidated", "score": None, "tier": None, "config_changed": False, "stale": False, "last_validated_at": None}

    score = meta.quality_score if meta else latest.get("score") if latest else None
    tier = meta.quality_tier if meta else None
    last_at = (meta.last_validated_at.isoformat() if meta and meta.last_validated_at else
               latest.get("created_at") if latest else None)

    # Check if workflow steps changed since last validation
    config_changed = False
    if latest:
        last_config = latest.get("extraction_config", {})
        current_steps = [{"name": s.get("name", ""), "tasks": s.get("tasks", [])} for s in (wf.get("steps_expanded", []) if isinstance(wf, dict) else [])]
        if not current_steps:
            # Fallback: hash validation_plan + step IDs
            current_config = {"validation_plan": wf.validation_plan if hasattr(wf, "validation_plan") else [], "steps": [str(s) for s in (wf.steps if hasattr(wf, "steps") else [])]}
        else:
            current_config = current_steps
        current_hash = hashlib.sha256(json.dumps(current_config, sort_keys=True, default=str).encode()).hexdigest()
        last_hash = hashlib.sha256(json.dumps(last_config, sort_keys=True, default=str).encode()).hexdigest()
        config_changed = current_hash != last_hash

    # Check staleness (>14 days)
    stale = False
    now_utc = _dt.datetime.now(_dt.timezone.utc)
    if meta and meta.last_validated_at:
        lv = meta.last_validated_at
        if lv.tzinfo is None:
            lv = lv.replace(tzinfo=_dt.timezone.utc)
        stale = (now_utc - lv).days > 14
    elif latest and latest.get("created_at"):
        from dateutil.parser import isoparse
        created = isoparse(latest["created_at"])
        if created.tzinfo is None:
            created = created.replace(tzinfo=_dt.timezone.utc)
        stale = (now_utc - created).days > 14

    return {
        "status": "validated",
        "score": score,
        "tier": tier,
        "last_validated_at": last_at,
        "config_changed": config_changed,
        "stale": stale,
    }


@router.get("/{workflow_id}/validation-plan", response_model=ValidationPlanResponse)
async def get_validation_plan(workflow_id: str, user: User = Depends(get_current_user)):
    try:
        checks = await svc.get_validation_plan(workflow_id, user=user)
        return ValidationPlanResponse(checks=checks)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))


@router.put("/{workflow_id}/validation-plan", response_model=ValidationPlanResponse)
async def update_validation_plan(
    workflow_id: str, req: UpdateValidationPlanRequest, user: User = Depends(get_current_user),
):
    try:
        checks = await svc.update_validation_plan(workflow_id, [c.model_dump() for c in req.checks], user=user)
        return ValidationPlanResponse(checks=checks)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))


@router.post("/{workflow_id}/validation-plan/generate", response_model=ValidationPlanResponse)
@limiter.limit("5/minute")
async def generate_validation_plan(request: Request, workflow_id: str, user: User = Depends(get_current_user)):
    try:
        checks = await svc.generate_validation_plan(workflow_id, user=user)
        return ValidationPlanResponse(checks=checks)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.get("/{workflow_id}/validation-inputs", response_model=ValidationInputsResponse)
async def get_validation_inputs(workflow_id: str, user: User = Depends(get_current_user)):
    try:
        inputs = await svc.get_validation_inputs(workflow_id, user=user)
        doc_uuids = [inp.get("document_uuid") for inp in inputs if inp.get("document_uuid")]
        exists_map = await _check_validation_input_documents_exist(doc_uuids)
        for inp in inputs:
            if inp.get("document_uuid"):
                inp["document_exists"] = exists_map.get(inp["document_uuid"], False)
        return ValidationInputsResponse(inputs=inputs)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))


@router.put("/{workflow_id}/validation-inputs", response_model=ValidationInputsResponse)
async def update_validation_inputs(
    workflow_id: str, req: UpdateValidationInputsRequest, user: User = Depends(get_current_user),
):
    try:
        inputs = await svc.update_validation_inputs(workflow_id, [i.model_dump() for i in req.inputs], user=user)
        return ValidationInputsResponse(inputs=inputs)
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))


@router.post("/{workflow_id}/create-temp-documents")
async def create_temp_documents(
    workflow_id: str, req: CreateTempDocumentsRequest, user: User = Depends(get_current_user),
):
    wf = await get_authorized_workflow(workflow_id, user)
    if not wf:
        raise HTTPException(status_code=404, detail="Workflow not found")
    uuids = await svc.create_temp_documents_from_text(req.texts, user.user_id)
    return {"document_uuids": uuids}


@router.post("/{workflow_id}/validate", response_model=ValidateWorkflowResponse)
async def validate_workflow(workflow_id: str, user: User = Depends(get_current_user)):
    try:
        result = await svc.validate_workflow(workflow_id, user=user)
        return ValidateWorkflowResponse(**result)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.post("/{workflow_id}/save-expected-output")
async def save_expected_output(workflow_id: str, request: Request, user: User = Depends(get_current_user)):
    """Mark a completed workflow execution as expected output for validation."""
    body = await request.json()
    session_id = body.get("session_id")
    label = body.get("label")
    if not session_id:
        raise HTTPException(status_code=400, detail="session_id is required")
    try:
        result = await svc.save_expected_output(workflow_id, session_id, user, label=label)
        return result
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.get("/{workflow_id}/expected-outputs")
async def get_expected_outputs(workflow_id: str, user: User = Depends(get_current_user)):
    """List stored expected outputs for a workflow."""
    try:
        outputs = await svc.get_expected_outputs(workflow_id, user)
        return {"expected_outputs": outputs}
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))


@router.delete("/{workflow_id}/expected-outputs/{expected_id}")
async def delete_expected_output(
    workflow_id: str, expected_id: str, user: User = Depends(get_current_user),
):
    """Remove a stored expected output."""
    ok = await svc.delete_expected_output(workflow_id, expected_id, user)
    if not ok:
        raise HTTPException(status_code=404, detail="Expected output not found")
    return {"ok": True}


# ---------------------------------------------------------------------------
# External API integration endpoints (x-api-key auth)
# ---------------------------------------------------------------------------


@router.post("/run-integrated")
@limiter.limit("10/minute")
async def run_workflow_integrated(
    request: Request,
    workflow_id: str = Form(...),
    files: list[UploadFile] = File(default=[]),
    user: User = Depends(get_api_key_user),
):
    """Run a workflow via external API with file uploads."""
    import uuid as _uuid
    from pathlib import Path
    from app.config import Settings
    from app.models.document import SmartDocument
    from app.tasks.upload_tasks import dispatch_upload_tasks
    from app.models.activity import ActivityType
    from app.services import activity_service

    settings = Settings()
    doc_uuids: list[str] = []

    for upload in files:
        if not upload.filename:
            continue
        uid = _uuid.uuid4().hex.upper()
        ext = (upload.filename.rsplit(".", 1)[-1] if "." in upload.filename else "pdf").lower()
        relative_path = Path(user.user_id) / f"{uid}.{ext}"
        upload_dir = Path(settings.upload_dir) / user.user_id
        upload_dir.mkdir(parents=True, exist_ok=True)
        file_path = upload_dir / f"{uid}.{ext}"
        file_data = await upload.read()
        file_path.write_bytes(file_data)

        doc = SmartDocument(
            title=upload.filename,
            processing=True,
            valid=True,
            raw_text="",
            downloadpath=str(relative_path),
            path=str(relative_path),
            extension=ext,
            uuid=uid,
            user_id=user.user_id,
            folder="0",
        )
        await doc.insert()

        task_id = dispatch_upload_tasks(
            document_uuid=uid, extension=ext, document_path=str(file_path),
            user_id=user.user_id,
        )
        doc.task_id = task_id
        await doc.save()
        doc_uuids.append(uid)

    if not doc_uuids:
        raise HTTPException(status_code=400, detail="No files provided")

    # Create activity
    activity = await activity_service.activity_start(
        type=ActivityType.WORKFLOW_RUN,
        title=f"API Workflow {workflow_id}",
        user_id=user.user_id,
        workflow=PydanticObjectId(workflow_id),
    )

    try:
        session_id = await svc.run_workflow(
            workflow_id, doc_uuids, user.user_id,
            user=user,
        )
        return {
            "status": "queued",
            "activity_id": str(activity.id),
            "session_id": session_id,
        }
    except ValueError as e:
        from app.models.activity import ActivityStatus
        await activity_service.activity_finish(activity.id, ActivityStatus.FAILED, error=str(e))
        raise HTTPException(status_code=404, detail=str(e))
